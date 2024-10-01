from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn


# PlaceholderReplacer 클래스 정의
class PlaceholderReplacer:
    """
    문서 내의 placeholder를 주어진 값으로 대체하는 클래스.

    Args:
        doc (Document): python-docx로 로드된 문서 객체.
        placeholder_values (dict): 대체할 placeholder와 해당 값을 담은 딕셔너리.
        font_name (str, optional): 대체된 텍스트에 적용할 글꼴 이름. 기본값은 None.
        font_size (int, optional): 대체된 텍스트에 적용할 글꼴 크기(pt 단위). 기본값은 None.
        font_color (tuple, optional): 대체된 텍스트에 적용할 글꼴 색상(RGB 값). 기본값은 None.
        
    """

    def __init__(self, doc, placeholder_values, font_name=None, font_size=None, font_color=None):
        """
        PlaceholderReplacer 클래스의 초기화 메서드.
        """
        self.doc = doc
        self.placeholder_values = placeholder_values
        self.font_name = font_name
        self.font_size = font_size
        self.font_color = font_color
        self.replaced_placeholders = {}


    def replace_placeholder(self, paragraph, placeholder, value, context):
        """
        문단 내에서 placeholder를 찾아 주어진 값으로 대체하는 메서드.

        Args:
            paragraph (Paragraph): 대체 작업을 진행할 문단 객체.
            placeholder (str): 대체할 placeholder.
            value (str): 대체할 값.
            context (str): 대체 작업이 이루어지는 위치(예: 문단 또는 테이블).

        Returns:
            bool: 대체가 성공하면 True, 실패하면 False.
            
        """
        if placeholder in paragraph.text:
            full_text = ""
            for run in paragraph.runs:  # 문단의 모든 run을 순회
                full_text += run.text
            new_text = full_text.replace(placeholder, value)

            # 모든 run을 지우고 새 텍스트를 단일 run에 추가
            paragraph.clear()
            run = paragraph.add_run(new_text)

            # 설정된 폰트 값이 있을 경우만 적용
            if self.font_name:
                run.font.name = self.font_name
                r = run._element
                r.rPr.rFonts.set(qn('w:eastAsia'), self.font_name)
            if self.font_size:
                run.font.size = Pt(self.font_size)
            if self.font_color:
                run.font.color.rgb = RGBColor(*self.font_color)

            # 주석처리된 로그 (필요 시 사용)
            # print(f"{context}에서 {placeholder} 대체 성공")
            return True
        return False


    def process_paragraphs(self):
        """
        문서의 모든 문단을 처리하여 placeholder를 대체하는 메서드.
        
        """
        for para in self.doc.paragraphs:  # 문서의 모든 문단을 순회
            for placeholder, value in self.placeholder_values.items():
                if placeholder not in self.replaced_placeholders:
                    if self.replace_placeholder(para, placeholder, value, context="문단"):
                        self.replaced_placeholders[placeholder] = "문단"


    def process_tables(self):
        """
        문서 내의 모든 테이블을 처리하여 placeholder를 대체하는 메서드.
        
        """
        for table in self.doc.tables:  # 문서의 모든 테이블을 순회
            for row in table.rows:  # 테이블의 모든 행을 순회
                for cell in row.cells:  # 행의 모든 셀을 순회
                    for para in cell.paragraphs:  # 셀 내의 모든 문단을 순회
                        for placeholder, value in self.placeholder_values.items():
                            if placeholder not in self.replaced_placeholders:
                                if self.replace_placeholder(para, placeholder, value, context="테이블"):
                                    self.replaced_placeholders[placeholder] = "테이블"


    def check_replacement_status(self):
        """
        placeholder 대체 상태를 확인하는 메서드. 성공 및 실패 개수를 반환.

        Returns:
            tuple: (성공한 placeholder 개수, 실패한 placeholder 개수)
            
        """
        success_count = 0
        fail_count = 0
        for placeholder in self.placeholder_values.keys():
            if placeholder in self.replaced_placeholders:
                location = self.replaced_placeholders[placeholder]
                # 주석처리된 로그 (필요 시 사용)
                # print(f"{placeholder}: {location}에서 성공")
                success_count += 1
            else:
                # 주석처리된 로그 (필요 시 사용)
                # print(f"{placeholder}: 실패")
                fail_count += 1
        # 주석처리된 로그 (필요 시 사용)
        # print(f"\n총 성공한 placeholder: {success_count}, 총 실패한 placeholder: {fail_count}")
        return success_count, fail_count



## 실제 사용 예시
if __name__ == '__main__':
    # 1. 문서 열기
    doc = Document('labeled-PCI-DSS-v4-0-SAQ-A-r1.docx')

    # 2. placeholder와 교체할 값 설정
    placeholder_values = {
        '{company_name}': 'marketian',
        '{dba}': 'hello, world',
        '{appendix_c1_title}': 'Requirements 3.5.2',
        '{appendix_c1_content}': 'Lorem ipsum, duaos einenskd enaoldfj',
        '{appendix_c2_title}': 'Requirements 3.5.2',
        '{appendix_c2_content}': 'Suspendisse ut purus sed quam consectetur cursus. Pellentesque eget metus tristique, tincidunt elit vel, posuere nulla. Aenean maximus purus eget mi consectetur, ac ultrices velit laoreet. Aliquam nulla urna, fermentum tempus interdum non, faucibus.',
    }

    # 3. PlaceholderReplacer 객체 생성
    replacer = PlaceholderReplacer(doc, placeholder_values)

    # 4. 문단과 테이블 내에서 placeholder 대체 수행
    replacer.process_paragraphs()
    replacer.process_tables()

    # 5. 대체 성공 여부 출력
    replacer.check_replacement_status()

    # 6. 변경된 문서 저장
    doc.save('output1.docx')

    # 주석처리된 로그를 사용하고 싶다면, 클래스 내 주석을 해제하세요.
