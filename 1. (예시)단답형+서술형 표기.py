import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn


def replace_placeholder(paragraph, placeholder, value, context, font_name=None, font_size=None, font_color=None):
    if placeholder in paragraph.text:
        full_text = ""
        for run in paragraph.runs:
            full_text += run.text
        new_text = full_text.replace(placeholder, value)
        
        # clear all runs and add the new text in a single run
        paragraph.clear()
        run = paragraph.add_run(new_text)
        
        # 설정된 폰트 값이 있을 경우만 적용
        if font_name:
            run.font.name = font_name
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        if font_size:
            run.font.size = Pt(font_size)
        if font_color:
            run.font.color.rgb = RGBColor(*font_color)
        return True
    return False


def process_paragraphs(doc, placeholder_values, replaced_placeholders, font_name, font_size, font_color):
    for para in doc.paragraphs:
        for placeholder, value in placeholder_values.items():
            if placeholder not in replaced_placeholders:  # 이미 대체된 경우 건너뛰기
                if replace_placeholder(para, placeholder, value, context="문단", font_name=font_name, font_size=font_size, font_color=font_color):
                    replaced_placeholders[placeholder] = '문단'


def process_tables(doc, placeholder_values, replaced_placeholders, font_name, font_size, font_color):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for placeholder, value in placeholder_values.items():
                        if placeholder not in replaced_placeholders:  # 이미 대체된 경우 건너뛰기
                            if replace_placeholder(para, placeholder, value, context="테이블", font_name=font_name, font_size=font_size, font_color=font_color):
                                replaced_placeholders[placeholder] = '테이블'


def check_replacement_status(placeholder_values, replaced_placeholders):
    success_count = 0
    fail_count = 0
    for placeholder in placeholder_values.keys():
        if placeholder in replaced_placeholders:
            location = replaced_placeholders[placeholder]
            print(f"{placeholder}: {location}에서 성공")
            success_count += 1
        else:
            print(f"{placeholder}: 실패")
            fail_count += 1
    print(f"\n총 성공한 placeholder: {success_count}, 총 실패한 placeholder: {fail_count}")


if __name__ == "__main__":
    # 시작 시간 기록
    start_time = time.time()

    # 템플릿 문서 열기
    doc = Document('labeled-PCI-DSS-v4-0-SAQ-A-r1.docx')

    # placeholder 대체하고 싶은 값들.
    placeholder_values = {
        '{company_name}': 'marketian',
        '{dba}': 'hello, world',
        '{appendix_c1_title}': 'Requirements 3.5.2',
        '{appendix_c1_content}': 'Lorem ipsum, duaos einenskd enaoldfj',
        '{appendix_c2_title}': 'Requirements 3.5.2',
        '{appendix_c2_content}': 'Suspendisse ut purus sed quam consectetur cursus. Pellentesque eget metus tristique, tincidunt elit vel, posuere nulla. Aenean maximus purus eget mi consectetur, ac ultrices velit laoreet. Aliquam nulla urna, fermentum tempus interdum non, faucibus.',
    }

    # 기본 폰트 설정 (None으로 설정)
    font_name = None
    font_size = None
    font_color = None

    # 대체된 placeholder 추적: 딕셔너리로 수정하여 위치 정보도 기록
    replaced_placeholders = {}

    # 일반 텍스트에서 placeholder 대체
    process_paragraphs(doc, placeholder_values, replaced_placeholders, font_name, font_size, font_color)

    # 테이블에서 placeholder 대체
    process_tables(doc, placeholder_values, replaced_placeholders, font_name, font_size, font_color)

    # 대체 성공 여부 체크 및 최종 출력
    check_replacement_status(placeholder_values, replaced_placeholders)

    # 결과 문서를 저장
    doc.save('output1.docx')

    # 총 소요 시간 계산 및 출력
    end_time = time.time()
    total_time = end_time - start_time
    print(f"\n총 작업 시간: {total_time:.4f} 초")
