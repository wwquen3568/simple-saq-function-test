from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn



def replace_placeholder(paragraph, placeholder, value, context, font_name=None, font_size=None, font_color=None):
    if placeholder in paragraph.text:
        print(f"{context}에서 {placeholder}를 대체했습니다.")
        full_text = ""
        for run in paragraph.runs:
            full_text += run.text
        new_text = full_text.replace(placeholder, value)
        
        # clear all runs and add the new text in a single run
        paragraph.clear()
        run = paragraph.add_run(new_text)
        
        # 설정된 폰트 값이 있을 경우만 적용
        if font_name:
            run.font.name = font_name
            # for Korean font name setting if needed
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        if font_size:
            run.font.size = Pt(font_size)
        if font_color:
            run.font.color.rgb = RGBColor(*font_color)


if __name__ == "__main__":
    # 템플릿 문서 열기
    doc = Document('labeled-PCI-DSS-v4-0-SAQ-A-r1.docx')

    # placeholder 대체하고 싶은 값들.
    placeholder_values = {
        ## 페이지 9, Section1: Assessment Information
        '{company_name}': 'marketian',
        '{dba}': 'hello, world',

        ## 페이지 34, Appendix C: Explanation of Requirements Noted ad Not Aplicable
        '{appendix_c1_title}': 'Requirements 3.5.2',
        '{appendix_c1_content}': 'Lorem ipsum, duaos einenskd enaoldfj',
        '{appendix_c2_title}': 'Requirements 3.5.2',
        '{appendix_c2_content}': 'Suspendisse ut purus sed quam consectetur cursus. Pellentesque eget metus tristique, tincidunt elit vel, posuere nulla. Aenean maximus purus eget mi consectetur, ac ultrices velit laoreet. Aliquam nulla urna, fermentum tempus interdum non, faucibus.',
        
        # 필요한 다른 placeholder와 값을 추가
    }

    # 기본 폰트 설정 (None으로 설정)
    font_name = None  # ex) "나눔스퀘어"
    font_size = None
    font_color = None

    # 경우 1) placeholder가 일반적으로 문서 페이지에 작성되어있는 경우
    # 문서의 모든 단락에서 placeholder 대체 (일반 텍스트 문서에 있는 placeholder)
    for para in doc.paragraphs:
        for placeholder, value in placeholder_values.items():
            replace_placeholder(para, placeholder, value, context="일반 텍스트", font_name=font_name, font_size=font_size, font_color=font_color)

    # 경우 2) placeholder가 표 안에 있는 경우
    # 문서의 모든 테이블 셀에서 placeholder 대체 (테이블에 있는 placeholder)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for placeholder, value in placeholder_values.items():
                        replace_placeholder(para, placeholder, value, context="테이블", font_name=font_name, font_size=font_size, font_color=font_color)

    # 결과 문서를 저장
    doc.save('output.docx')
    print("작업 완료")
